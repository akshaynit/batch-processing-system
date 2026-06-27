"""Assemble the design document (.docx) from text content + rendered diagrams.

This is intentionally data-driven so the document can be regenerated/updated
by editing the content below and re-running:

    python scripts/build_doc.py

Diagrams are expected as PNGs in docs/assets (produced by render_diagrams.py).
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
ASSETS = ROOT / "docs" / "assets"
OUT = ROOT / "docs" / "Batch_Eval_Engine_Design.docx"

ACCENT = RGBColor(0x1F, 0x4E, 0x79)
MUTED = RGBColor(0x59, 0x59, 0x59)


# --------------------------------------------------------------------------- #
# Helpers
# --------------------------------------------------------------------------- #
def set_orientation(section, landscape: bool) -> None:
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        if section.page_width < section.page_height:
            section.page_width, section.page_height = section.page_height, section.page_width
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        if section.page_width > section.page_height:
            section.page_width, section.page_height = section.page_height, section.page_width


def caption(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = MUTED
    p.paragraph_format.space_after = Pt(10)


def image(doc, name: str, width_in: float, cap: str | None = None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(ASSETS / name), width=Inches(width_in))
    if cap:
        caption(doc, cap)


def landscape_image(doc, name: str, heading: str, width_in: float, cap: str) -> None:
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    set_orientation(sec, landscape=True)
    doc.add_heading(heading, level=2)
    image(doc, name, width_in, cap)
    sec2 = doc.add_section(WD_SECTION.NEW_PAGE)
    set_orientation(sec2, landscape=False)


def body(doc, text: str, bold: bool = False, space_after: int = 6) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    p.paragraph_format.space_after = Pt(space_after)


def bullet(doc, text: str, sub: bool = False) -> None:
    style = "List Bullet 2" if sub else "List Bullet"
    try:
        p = doc.add_paragraph(style=style)
    except KeyError:
        p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)


def monospace(doc, text: str, size: float = 8.5) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    lines = text.split("\n")
    for i, line in enumerate(lines):
        r = p.add_run(line if line else " ")
        r.font.name = "Consolas"
        r.font.size = Pt(size)
        if i != len(lines) - 1:
            r.add_break()


def table(doc, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        t.style = "Light Grid Accent 1"
    except KeyError:
        t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(htext)
        run.bold = True
        run.font.size = Pt(9.5)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(val)
            run.font.size = Pt(9.5)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def h1(doc, text: str) -> None:
    doc.add_heading(text, level=1)


def h2(doc, text: str) -> None:
    doc.add_heading(text, level=2)


# --------------------------------------------------------------------------- #
# Build
# --------------------------------------------------------------------------- #
def build() -> None:
    doc = Document()

    # Base styling
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    for lvl, sz, col in [("Heading 1", 18, ACCENT), ("Heading 2", 14, ACCENT), ("Heading 3", 12, ACCENT)]:
        st = doc.styles[lvl]
        st.font.color.rgb = col
        st.font.size = Pt(sz)

    # ---- Cover ----
    for _ in range(4):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Asynchronous Batch Evaluation Engine")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = ACCENT
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Architecture & Low-Level Design")
    r.font.size = Pt(16)
    r.font.color.rgb = MUTED
    for _ in range(2):
        doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run("Version 0.1 (Draft)  ·  Living document")
    r.font.size = Pt(11)
    r.italic = True
    doc.add_page_break()

    # ---- 1. Purpose ----
    h1(doc, "1. Purpose & Status")
    body(
        doc,
        "This document describes the architecture and low-level design of a production-ready "
        "asynchronous batch evaluation engine. It is intended as a presentation and discussion "
        "artifact: it explains the system at the level of components, data flow, and class design "
        "without requiring the reader to look at source code.",
    )
    body(doc, "It is a living document and will be updated iteratively until the final version.", bold=True)

    # ---- 2. System Overview ----
    h1(doc, "2. System Overview")
    body(
        doc,
        "The engine reads a local file containing an array of prompts, immediately returns an "
        "operational job id, and executes the batch fully in the background. It partitions the "
        "prompts into chunks and fans them out across a bounded worker pool that invokes a live "
        "inference endpoint. It enforces upstream rate-limit backpressure with exponential back-off "
        "and jitter, isolates individual row failures, and aggregates results into structured, "
        "downloadable reports.",
    )
    body(doc, "Progress and results are exposed via a small HTTP API.")
    bullet(doc, "POST /jobs  -  ingest a batch file; returns a job id immediately (non-blocking).")
    bullet(doc, "GET /job/{id}/status  -  live progress: total / succeeded / failed / in-progress.")
    bullet(doc, "GET /job/{id}/download  -  the final compiled results array (success + failure details).")

    # ---- 3. Functional Requirements ----
    h1(doc, "3. Functional Requirements")
    table(
        doc,
        ["#", "Requirement", "How it is met"],
        [
            ["1", "Batch file ingestion (non-blocking, returns job id)",
             "Streaming parse of sample_batch.json; job id returned instantly; ingestion runs in background."],
            ["2", "Scatter-gather distribution across a bounded worker pool",
             "Rows partitioned into chunks; workers pull chunks from Postgres and call the live endpoint."],
            ["3", "API backpressure control",
             "Concurrency semaphore + exponential back-off with full jitter; no element is dropped."],
            ["4", "Partial failure & reporting",
             "Each row failure isolated; results compiled; tracked via /status and retrieved via /download."],
        ],
        widths=[0.4, 2.7, 3.4],
    )

    # ---- 4. Key Design Decisions ----
    h1(doc, "4. Key Design Decisions")
    table(
        doc,
        ["Area", "Choice", "Rationale"],
        [
            ["Language / API", "Python + FastAPI",
             "Async-native fit for high-concurrency outbound I/O; auto OpenAPI docs; Pydantic validation."],
            ["Inference provider", "DigitalOcean Serverless Inference",
             "OpenAI-compatible per-request endpoint; pay-per-token; real rate limits exercise backpressure."],
            ["Model", "llama3-8b-instruct",
             "Cheapest/fastest 8B-class instruct model; matches the cost-efficient, high-throughput goal."],
            ["Queue & state", "Postgres (FOR UPDATE SKIP LOCKED)",
             "Durable multi-worker claim queue with crash recovery; one system for storage + queueing."],
            ["Result storage", "Pluggable StorageBackend (local now, DO Spaces later)",
             "S3-compatible abstraction; switch to Spaces by config only, no logic change."],
            ["Concurrency model", "Bounded worker pool + semaphore",
             "Caps in-flight calls to respect rate limits; constant memory regardless of dataset size."],
        ],
        widths=[1.3, 2.2, 3.0],
    )

    # ---- 5. Architecture ----
    h1(doc, "5. Architecture")
    h2(doc, "5.1 Architecture Overview")
    image(doc, "01_architecture_flow.png", 6.6,
          "Figure 1. End-to-end architecture: the four tracks and the shared datastores.")
    body(doc, "The pipeline is organised into four tracks:")
    bullet(doc, "1 - Input Ingestion: POST /jobs creates the job, returns the id immediately, and a "
                "background task stream-parses and validates the file, inserting PENDING rows.")
    bullet(doc, "2 - Scatter / Segmentation: a bounded worker pool pulls chunks of rows from the "
                "Postgres queue using FOR UPDATE SKIP LOCKED.")
    bullet(doc, "3 - Backpressure / Throttling: workers call the live endpoint under a concurrency "
                "semaphore with exponential back-off + jitter; permanent failures are isolated per row.")
    bullet(doc, "4 - Gather / Collection: completed chunks are flushed to the storage backend and "
                "marked done; /status reports counts and /download assembles the final array.")

    h2(doc, "5.2 Worker Processing Detail")
    image(doc, "02_worker_detail.png", 4.0,
          "Figure 2. Per-chunk worker processing, retry loop, checkpointing, and lease heartbeat.")

    h2(doc, "5.3 Backpressure & Throttling")
    body(doc, "Backpressure is the ability to automatically slow intake when the inference endpoint is "
              "saturated, and speed back up when it recovers - without dropping work. It is built from "
              "three cooperating mechanisms, each a separately tunable knob:")
    table(
        doc,
        ["Mechanism", "Limits", "Prevents", "Knob"],
        [
            ["Semaphore", "concurrent in-flight calls", "429 storms, connection/memory blow-up", "MAX_CONCURRENCY"],
            ["Bounded pool", "chunks processed in parallel", "unbounded prefetch into memory", "WORKER_POOL_SIZE"],
            ["Back-off + jitter", "retry timing", "hammering the API, thundering herd", "BASE/CAP/MAX_ATTEMPTS"],
            ["Lease + heartbeat", "(robustness)", "losing work if a worker dies mid-call", "LEASE_TTL_SECONDS"],
        ],
        widths=[1.4, 1.8, 2.4, 1.4],
    )
    image(doc, "03_backpressure_sequence.png", 4.6,
          "Figure 3. Adaptive throttling: permits held during back-off naturally reduce request rate.")
    body(doc, "Because a retry sleeps while still holding its semaphore permit, a burst of 429s causes "
              "fewer new calls to start, so the request rate to the provider drops on its own until it "
              "recovers - the system self-regulates around the provider's limit.")

    h2(doc, "5.4 Work-Item Lifecycle & Concurrency Safety")
    image(doc, "05_item_state.png", 4.2,
          "Figure 4. State machine for each prompt row.")
    body(doc, "Postgres is the single source of truth for state. Every transition is an atomic, guarded "
              "SQL update. The claim/lease/guard trio removes all multi-worker races:")
    table(
        doc,
        ["Race", "How it is eliminated"],
        [
            ["Two workers claim the same row", "FOR UPDATE SKIP LOCKED - the second transaction skips locked rows."],
            ["Worker crashes mid-chunk", "Lease (visibility timeout): expired CLAIMED rows are re-claimed automatically."],
            ["Zombie writer after reassignment", "Completion update guarded by worker_id; stale update affects 0 rows."],
            ["Duplicate ingestion on restart", "UNIQUE (job_id, external_id) + ON CONFLICT DO NOTHING."],
            ["Duplicate result object", "Deterministic chunk key => idempotent overwrite (effectively exactly-once)."],
        ],
        widths=[2.2, 4.3],
    )

    # ---- 6. Memory & Scaling ----
    h1(doc, "6. Memory Footprint & Scaling to 500,000 Items")
    body(doc, "Core principle: memory must be O(concurrency x chunk_size), not O(dataset_size). If memory "
              "grows with the number of items, the process eventually crashes (OOM). Each stage is kept bounded:")
    table(
        doc,
        ["Stage", "Naive (OOM risk)", "Our approach (bounded)"],
        [
            ["Ingestion", "json.load entire file (~2-5 GB at 500k)", "Stream-parse (ijson) + batched inserts; peak ~ one batch."],
            ["Distribution", "load all ids into memory", "Workers pull chunks on demand from Postgres."],
            ["In-flight", "fire all 500k at once", "Semaphore caps concurrency; bounded prefetch."],
            ["Results", "append all to one list", "Flush each chunk to storage, then free it."],
            ["/download", "build full array in memory", "Streaming response reads chunk objects sequentially."],
            ["/status", "load rows to count", "SQL aggregate (GROUP BY status) - O(1) memory."],
        ],
        widths=[1.1, 2.6, 2.8],
    )
    body(doc, "Memory ceiling (every term is a tunable constant - none depends on 500k):")
    monospace(doc,
              "peak_RAM  ~=  ingest_batch\n"
              "            +  (MAX_CONCURRENCY x avg_request_buffer)\n"
              "            +  (WORKER_POOL_SIZE x chunk_size x avg_item_size)\n"
              "            +  DB / HTTP connection-pool overhead")
    body(doc, "With MAX_CONCURRENCY=10, CHUNK_SIZE=50, ~2 KB/item, resident memory is single-digit to "
              "low-tens of MB whether the dataset is 1k, 500k, or 5M. Dataset size only changes runtime "
              "and the number of DB rows / storage objects - both on disk, not in RAM.")
    body(doc, "The real ceilings at scale (not memory):")
    bullet(doc, "Provider rate limit - the true throughput ceiling; scale horizontally (more workers / keys), "
                "not by enlarging the in-process pool.")
    bullet(doc, "Postgres queue contention - handled by SKIP LOCKED + an index on (job_id, status, lease_until).")
    bullet(doc, "Storage object count - tune chunk_size to balance object count vs crash-rework granularity.")

    # ---- 7. Low-Level Design ----
    h1(doc, "7. Low-Level Design")

    h2(doc, "7.1 Design Principles (SOLID)")
    table(
        doc,
        ["Principle", "Applied as"],
        [
            ["SRP", "Each class has one job (Worker processes chunks; RetryPolicy only computes retry decisions)."],
            ["OCP", "New provider/storage = a new class implementing an interface; engine unchanged."],
            ["LSP", "Mock and concrete adapters are fully substitutable behind ABCs (enables tests)."],
            ["ISP", "Small focused ports: InferenceClient, StorageBackend, JobRepository, WorkQueueRepository."],
            ["DIP", "Core depends on abstractions injected via constructors; concretes wired in main.py only."],
        ],
        widths=[1.0, 5.5],
    )

    h2(doc, "7.2 Project Structure")
    monospace(doc, PROJECT_TREE, size=8.5)

    h2(doc, "7.4 Database Schema")
    body(doc, "jobs - one row per submitted batch:")
    table(
        doc,
        ["Column", "Type", "Notes"],
        [
            ["id", "UUID (PK)", "Job identifier returned to the client."],
            ["state", "enum", "PENDING / INGESTING / RUNNING / COMPLETED / FAILED."],
            ["input_path", "text", "Workspace path of the batch file."],
            ["model", "text", "Inference model id."],
            ["total_items", "int", "Count after ingestion."],
            ["config", "jsonb", "Per-job overrides (concurrency, chunk size, etc.)."],
            ["created_at / updated_at", "timestamptz", "Timestamps."],
        ],
        widths=[1.8, 1.4, 3.3],
    )
    body(doc, "prompt_items - the durable work queue (one row per prompt):")
    table(
        doc,
        ["Column", "Type", "Notes"],
        [
            ["id", "bigserial (PK)", "Internal row id."],
            ["job_id", "UUID (FK)", "Owning job."],
            ["external_id", "text", "Id from the input file; UNIQUE per job."],
            ["prompt", "text", "Prompt text."],
            ["status", "enum", "PENDING / CLAIMED / SUCCEEDED / FAILED."],
            ["attempts", "smallint", "Retry counter."],
            ["worker_id / lease_until", "text / timestamptz", "Claim ownership + visibility timeout."],
            ["chunk_id / result_ref", "text", "Checkpoint chunk id and storage object key."],
            ["error", "text", "Reason if FAILED."],
        ],
        widths=[1.9, 1.7, 2.9],
    )
    body(doc, "Indexes: UNIQUE (job_id, external_id) for dedupe/idempotent ingest; partial index on "
              "(job_id, status, lease_until) for fast claims. Claims use the atomic "
              "FOR UPDATE SKIP LOCKED pattern inside short transactions.")

    h2(doc, "7.5 Configuration")
    table(
        doc,
        ["Setting", "Default", "Purpose"],
        [
            ["LLM_BASE_URL", "https://inference.do-ai.run/v1/", "DO serverless endpoint (OpenAI-compatible)."],
            ["LLM_MODEL", "llama3-8b-instruct", "Cheap/fast inference model."],
            ["WORKER_POOL_SIZE", "4", "Chunks processed in parallel."],
            ["MAX_CONCURRENCY", "10", "Max simultaneous inference calls (semaphore)."],
            ["CHUNK_SIZE", "50", "Rows per claim + checkpoint granularity."],
            ["RETRY_MAX_ATTEMPTS", "5", "Max attempts before a row is FAILED."],
            ["RETRY_BASE / FACTOR / CAP", "1.0 / 2.0 / 30.0", "Exponential back-off with full jitter."],
            ["LEASE_TTL_SECONDS", "120", "Visibility timeout for crash recovery."],
            ["STORAGE_BACKEND", "local", "local now; s3 (DO Spaces) later."],
        ],
        widths=[2.2, 1.9, 2.4],
    )

    h2(doc, "7.6 Validation Rules")
    body(doc, "Batch-fatal (reject at POST /jobs - nothing created):")
    bullet(doc, "Unknown/forbidden model, or parameters out of range.")
    bullet(doc, "Projected cost over budget, or item count over MAX_ITEMS_PER_BATCH.")
    bullet(doc, "Duplicate ids (when policy = reject); unreadable / non-array file.")
    body(doc, "Row-isolated (recorded as FAILED, job continues - per requirement 4):")
    bullet(doc, "Empty / oversized prompt; prompt tokens + max_tokens exceed context window.")
    bullet(doc, "Unfilled template placeholders; encoding issues; per-row parameter errors.")

    # ---- Landscape: component + class diagrams ----
    landscape_image(doc, "04_component_diagram.png", "7.3 Component & Dependency View", 9.2,
                    "Figure 5. Components depend on interfaces (ports); infrastructure provides adapters; "
                    "main.py is the only composition root. Arrows point toward abstractions (DIP).")
    landscape_image(doc, "06_class_uml.png", "7.7 Class Model (UML)", 9.4,
                    "Figure 6. Core classes, interfaces, and their relationships.")
    table(
        doc,
        ["Class / Interface", "Responsibility"],
        [
            ["JobService", "Orchestrates submit / status / download; depends only on interfaces."],
            ["BatchIngestor", "Streams the file, validates rows, bulk-inserts into the queue."],
            ["ValidationPipeline", "Composite of validation rules (batch-fatal + row-isolated)."],
            ["BatchEngine", "Spawns and supervises the bounded worker pool; finalizes job state."],
            ["Worker", "Claims a chunk, processes rows (retry+limiter), flushes results, heartbeats lease."],
            ["RetryPolicy", "Computes retry eligibility and back-off + jitter delay."],
            ["ConcurrencyLimiter", "Semaphore enforcing the hard ceiling on in-flight calls."],
            ["InferenceClient (port)", "Live endpoint contract; DOInferenceClient + MockInferenceClient."],
            ["StorageBackend (port)", "Result persistence; LocalFileStorage + S3SpacesStorage."],
            ["JobRepository / WorkQueueRepository (ports)", "Data access; Postgres adapters implement the claim/complete SQL."],
        ],
        widths=[2.4, 4.1],
    )

    # ---- 8. Testing ----
    h1(doc, "8. Testing Strategy")
    body(doc, "Dependency inversion makes the logic testable without a database or live API.")
    h2(doc, "Unit tests (fast, mocked)")
    bullet(doc, "RetryPolicy: delay bounds, jitter range, retryable mapping.")
    bullet(doc, "ValidationPipeline and each rule: batch-fatal vs row-isolated outcomes.")
    bullet(doc, "LocalFileStorage round-trip; ConcurrencyLimiter never exceeds N in-flight.")
    bullet(doc, "Worker with MockInferenceClient: success, retry-then-succeed, permanent-fail isolation.")
    bullet(doc, "BatchIngestor with a fixture file including invalid rows.")
    h2(doc, "Integration tests (real Postgres)")
    bullet(doc, "Repository claim/complete; concurrency race test (no item claimed twice).")
    bullet(doc, "Lease expiry / reclaim; idempotent ingest.")
    bullet(doc, "Full end-to-end via the API with the mock client: submit -> status -> download.")

    # ---- 9. CI/CD ----
    h1(doc, "9. CI/CD Pipeline")
    body(doc, "GitHub Actions on push and pull_request:")
    bullet(doc, "Lint (ruff) and type-check (mypy).")
    bullet(doc, "Unit tests.")
    bullet(doc, "Integration tests against a Postgres service container (alembic upgrade head first).")
    bullet(doc, "Coverage upload + threshold gate. No live API keys in CI - everything uses the mock client.")

    # ---- 10. Docs ----
    h1(doc, "10. Documentation Deliverables")
    bullet(doc, "README - setup, environment, docker-compose, run instructions, curl examples.")
    bullet(doc, "This design document (architecture + LLD), kept up to date.")
    bullet(doc, "Auto-generated OpenAPI / Swagger UI at /docs from FastAPI.")

    # ---- Change log ----
    h1(doc, "11. Change Log")
    table(
        doc,
        ["Version", "Notes"],
        [["0.1", "Initial draft: architecture flows, memory/scaling analysis, and low-level design."]],
        widths=[1.2, 5.3],
    )

    doc.save(OUT)
    print(f"Wrote {OUT}")


PROJECT_TREE = """batch-eval-engine/
  app/
    main.py                  # FastAPI app factory + composition root (DI wiring)
    api/
      dependencies.py        # provider funcs (inject services into routes)
      routers/jobs.py        # POST /jobs, GET /job/{id}/status, /download
    core/
      config.py              # Settings (pydantic-settings)
      exceptions.py          # typed domain exceptions
    domain/
      enums.py               # JobState, ItemStatus
      models.py              # PromptItem, ClaimedItem, JobStatus, ...
      results.py             # InferenceRequest/Response, ResultRecord
    interfaces/              # PORTS (ABCs) - the contracts
      inference_client.py
      storage_backend.py
      repositories.py
      validation.py
    infrastructure/         # ADAPTERS (implementations)
      db/ (session, ORM models, migrations, repositories/)
      inference/ (openai_client.py, mock_client.py)
      storage/ (local_storage.py, s3_storage.py)
    services/
      job_service.py
      ingestor.py
      validation/ (pipeline.py, rules.py)
    engine/
      batch_engine.py        # supervises the bounded worker pool
      worker.py              # claim -> process -> flush -> complete
      concurrency.py         # ConcurrencyLimiter (semaphore)
      retry.py               # RetryPolicy (back-off + jitter)
    utils/ (tokens.py, cost.py)
  workers/run_worker.py      # standalone worker-process entrypoint
  scripts/generate_sample_batch.py
  tests/ (unit/, integration/, conftest.py)
  docs/ (this document, diagrams/)
  .github/workflows/ci.yml
  docker-compose.yml         # Postgres (+ MinIO optional)
  alembic.ini
  .env.example
  pyproject.toml
  README.md"""


if __name__ == "__main__":
    build()
